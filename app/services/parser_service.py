"""
Production-ready document parsing service for PDF and DOCX files.
Supports comprehensive parsing with metadata extraction, chunking, and error handling.
"""

import logging
import traceback
import hashlib
import mimetypes
from typing import Dict, List, Any, Optional, Tuple
from dataclasses import dataclass
from pathlib import Path
import tempfile
import os
import re

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    logging.warning("PyMuPDF not available. PDF parsing will be disabled.")

try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    logging.warning("python-docx not available. DOCX parsing will be disabled.")

from app.core.config import settings

logger = logging.getLogger(__name__)


@dataclass
class DocumentChunk:
    """Represents a chunk of document text with metadata."""
    content: str
    page_number: int
    chunk_index: int
    start_position: int
    end_position: int
    metadata: Optional[Dict[str, Any]] = None


@dataclass
class DocumentMetadata:
    """Document metadata extracted during parsing."""
    title: Optional[str] = None
    author: Optional[str] = None
    subject: Optional[str] = None
    creator: Optional[str] = None
    creation_date: Optional[str] = None
    modification_date: Optional[str] = None
    page_count: int = 0
    word_count: int = 0
    character_count: int = 0
    file_size: int = 0
    file_hash: Optional[str] = None


@dataclass
class ParsedDocument:
    """Complete parsed document with all content and metadata."""
    content: str
    chunks: List[DocumentChunk]
    metadata: DocumentMetadata
    sections: List[Dict[str, Any]]
    parsing_stats: Dict[str, Any]
    
    def get_memory_estimate_mb(self) -> float:
        """Estimate memory usage in MB."""
        try:
            size = len(self.content.encode('utf-8'))
            size += len(str(self.metadata.__dict__).encode('utf-8'))
            size += len(str(self.sections).encode('utf-8'))
            size += len(str(self.parsing_stats).encode('utf-8'))
            if self.chunks:
                size += sum(len(str(chunk.__dict__).encode('utf-8')) for chunk in self.chunks)
            return size / (1024 * 1024)
        except Exception:
            return 0.0


class DocumentParsingError(Exception):
    """Raised when document parsing fails."""
    pass


class FileSizeError(DocumentParsingError):
    """Raised when file exceeds size limits."""
    pass


class PageLimitError(DocumentParsingError):
    """Raised when document exceeds page limits."""
    pass


class UnsupportedFormatError(DocumentParsingError):
    """Raised for unsupported file formats."""
    pass


class DocumentParserService:
    """
    Production-ready document parsing service.
    
    Features:
    - PDF parsing with PyMuPDF (metadata, text extraction, page-wise processing)
    - DOCX parsing with python-docx (structured content, paragraphs, tables)
    - Intelligent text chunking for optimal processing
    - Comprehensive error handling and validation
    - Memory-efficient processing for large documents
    - File integrity validation and security checks
    """
    
    def __init__(self):
        self.max_file_size_bytes = settings.MAX_FILE_SIZE_MB * 1024 * 1024
        self.max_pages = settings.MAX_PAGES
        self.allowed_extensions = [ext.lower() for ext in settings.ALLOWED_EXTENSIONS]
        self.chunk_size = 1000  # Default chunk size in characters
        self.chunk_overlap = 200  # Overlap between chunks
        
        # Validate dependencies
        self._validate_dependencies()
    
    def _validate_dependencies(self):
        """Validate that required parsing libraries are available."""
        if not PYMUPDF_AVAILABLE and not PYTHON_DOCX_AVAILABLE:
            raise RuntimeError("No document parsing libraries available. Install pymupdf and python-docx.")
        
        if not PYMUPDF_AVAILABLE:
            logger.warning("PyMuPDF not available - PDF parsing disabled")
        
        if not PYTHON_DOCX_AVAILABLE:
            logger.warning("python-docx not available - DOCX parsing disabled")
    
    async def parse_document(
        self, 
        file_content: bytes, 
        filename: str,
        validate_limits: bool = True
    ) -> ParsedDocument:
        """
        Parse a document from raw bytes.
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            validate_limits: Whether to enforce file size and page limits
            
        Returns:
            ParsedDocument with all extracted content and metadata
            
        Raises:
            DocumentParsingError: For various parsing failures
            FileSizeError: If file exceeds size limits
            PageLimitError: If document exceeds page limits
            UnsupportedFormatError: For unsupported file types
        """
        
        try:
            logger.info(f"Starting document parsing: {filename} ({len(file_content)} bytes)")
            
            # Validate file
            await self._validate_file(file_content, filename, validate_limits)
            
            # Determine file type and parse accordingly
            file_extension = Path(filename).suffix.lower()
            
            if file_extension == '.pdf':
                if not PYMUPDF_AVAILABLE:
                    raise UnsupportedFormatError("PDF parsing not available - PyMuPDF not installed")
                parsed_doc = await self._parse_pdf(file_content, filename)
                
            elif file_extension == '.docx':
                if not PYTHON_DOCX_AVAILABLE:
                    raise UnsupportedFormatError("DOCX parsing not available - python-docx not installed")
                parsed_doc = await self._parse_docx(file_content, filename)
                
            else:
                raise UnsupportedFormatError(f"Unsupported file format: {file_extension}")
            
            # Validate page limits after parsing
            if validate_limits and parsed_doc.metadata.page_count > self.max_pages:
                raise PageLimitError(
                    f"Document has {parsed_doc.metadata.page_count} pages, "
                    f"maximum allowed is {self.max_pages}"
                )
            
            # Add file hash for integrity verification
            parsed_doc.metadata.file_hash = self._calculate_file_hash(file_content)
            parsed_doc.metadata.file_size = len(file_content)
            
            logger.info(
                f"Document parsing completed: {filename} - "
                f"{parsed_doc.metadata.page_count} pages, "
                f"{parsed_doc.metadata.word_count} words, "
                f"{len(parsed_doc.chunks)} chunks"
            )
            
            return parsed_doc
            
        except (DocumentParsingError, FileSizeError, PageLimitError, UnsupportedFormatError):
            # Re-raise known errors
            raise
            
        except Exception as e:
            error_msg = f"Unexpected error parsing document {filename}: {str(e)}"
            logger.error(f"{error_msg}\n{traceback.format_exc()}")
            raise DocumentParsingError(error_msg) from e
    
    async def _validate_file(self, file_content: bytes, filename: str, validate_limits: bool):
        """Validate file before parsing."""
        
        # Check file size
        if validate_limits and len(file_content) > self.max_file_size_bytes:
            raise FileSizeError(
                f"File size {len(file_content)} bytes exceeds maximum "
                f"allowed size of {self.max_file_size_bytes} bytes"
            )
        
        # Check file extension
        file_extension = Path(filename).suffix.lower()
        if file_extension not in self.allowed_extensions:
            raise UnsupportedFormatError(
                f"File extension {file_extension} not supported. "
                f"Allowed extensions: {self.allowed_extensions}"
            )
        
        # Basic file content validation
        if len(file_content) == 0:
            raise DocumentParsingError("Empty file provided")
        
        # Check for potential security issues
        await self._security_validate_file(file_content, filename)
    
    async def _security_validate_file(self, file_content: bytes, filename: str):
        """Perform basic security validation on file content."""
        
        # Check for suspicious file names
        if '..' in filename or '/' in filename or '\\' in filename:
            raise DocumentParsingError("Filename contains invalid characters")
        
        # Basic file signature validation
        file_extension = Path(filename).suffix.lower()
        
        if file_extension == '.pdf':
            # PDF files should start with %PDF
            if not file_content.startswith(b'%PDF'):
                raise UnsupportedFormatError("Invalid PDF file signature")
                
        elif file_extension == '.docx':
            # DOCX files are ZIP archives, should start with PK
            if not file_content.startswith(b'PK'):
                raise UnsupportedFormatError("Invalid DOCX file signature")
    
    async def _parse_pdf(self, file_content: bytes, filename: str) -> ParsedDocument:
        """Parse PDF document using PyMuPDF."""
        
        try:
            # Create temporary file for PyMuPDF
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
            
            try:
                # Open PDF with PyMuPDF
                pdf_document = fitz.open(temp_file_path)
                
                # Extract metadata
                metadata = self._extract_pdf_metadata(pdf_document, filename)
                
                # Extract text content page by page
                full_content = ""
                chunks = []
                sections = []
                
                # Store page count before closing document
                total_pages = len(pdf_document)
                
                for page_num in range(total_pages):
                    page = pdf_document[page_num]
                    page_text = page.get_text()
                    
                    if page_text.strip():
                        # Add page text to full content
                        full_content += f"\n\n--- Page {page_num + 1} ---\n"
                        page_start_pos = len(full_content)
                        full_content += page_text
                        page_end_pos = len(full_content)
                        
                        # Create page-level chunks
                        page_chunks = self._create_chunks(
                            page_text, 
                            page_num + 1, 
                            len(chunks),
                            page_start_pos
                        )
                        chunks.extend(page_chunks)
                        
                        # Create section info for page
                        sections.append({
                            "type": "page",
                            "number": page_num + 1,
                            "title": f"Page {page_num + 1}",
                            "start_pos": page_start_pos,
                            "end_pos": page_end_pos,
                            "word_count": len(page_text.split())
                        })
                
                pdf_document.close()
                
                # Update metadata with content statistics
                metadata.word_count = len(full_content.split())
                metadata.character_count = len(full_content)
                
                # Parsing statistics
                parsing_stats = {
                    "parser_used": "PyMuPDF",
                    "total_pages": total_pages,
                    "pages_with_text": len([s for s in sections if s.get("word_count", 0) > 0]),
                    "total_chunks": len(chunks),
                    "average_chunk_size": sum(len(chunk.content) for chunk in chunks) / len(chunks) if chunks else 0
                }
                
                return ParsedDocument(
                    content=full_content,
                    chunks=chunks,
                    metadata=metadata,
                    sections=sections,
                    parsing_stats=parsing_stats
                )
                
            finally:
                # Clean up temporary file
                os.unlink(temp_file_path)
                
        except Exception as e:
            error_msg = f"PDF parsing failed for {filename}: {str(e)}"
            logger.error(f"{error_msg}\n{traceback.format_exc()}")
            raise DocumentParsingError(error_msg) from e
    
    async def _parse_docx(self, file_content: bytes, filename: str) -> ParsedDocument:
        """Parse DOCX document using python-docx."""
        
        try:
            # Create temporary file for python-docx
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
            
            try:
                # Open DOCX with python-docx
                document = Document(temp_file_path)
                
                # Extract metadata
                metadata = self._extract_docx_metadata(document, filename)
                
                # Extract text content
                full_content = ""
                chunks = []
                sections = []
                section_num = 0
                
                # Process paragraphs
                for para_idx, paragraph in enumerate(document.paragraphs):
                    para_text = paragraph.text.strip()
                    
                    if para_text:
                        # Check if this is a heading (basic heuristic)
                        is_heading = (
                            len(para_text) < 100 and 
                            (para_text.isupper() or 
                             para_text.title() == para_text or
                             paragraph.style.name.startswith('Heading'))
                        )
                        
                        if is_heading:
                            section_num += 1
                            sections.append({
                                "type": "heading",
                                "number": section_num,
                                "title": para_text,
                                "start_pos": len(full_content),
                                "end_pos": len(full_content) + len(para_text),
                                "word_count": len(para_text.split())
                            })
                        
                        # Add paragraph to full content
                        para_start_pos = len(full_content)
                        full_content += para_text + "\n\n"
                        para_end_pos = len(full_content)
                        
                        # Create chunks for paragraph
                        para_chunks = self._create_chunks(
                            para_text, 
                            1,  # DOCX doesn't have clear page numbers
                            len(chunks),
                            para_start_pos
                        )
                        chunks.extend(para_chunks)
                
                # Process tables
                for table_idx, table in enumerate(document.tables):
                    table_text = self._extract_table_text(table)
                    if table_text.strip():
                        table_start_pos = len(full_content)
                        full_content += f"\n\n--- Table {table_idx + 1} ---\n"
                        full_content += table_text + "\n\n"
                        table_end_pos = len(full_content)
                        
                        sections.append({
                            "type": "table",
                            "number": table_idx + 1,
                            "title": f"Table {table_idx + 1}",
                            "start_pos": table_start_pos,
                            "end_pos": table_end_pos,
                            "word_count": len(table_text.split())
                        })
                        
                        # Create chunks for table
                        table_chunks = self._create_chunks(
                            table_text, 
                            1,
                            len(chunks),
                            table_start_pos
                        )
                        chunks.extend(table_chunks)
                
                # Update metadata with content statistics
                metadata.page_count = 1  # DOCX doesn't have clear page concept
                metadata.word_count = len(full_content.split())
                metadata.character_count = len(full_content)
                
                # Parsing statistics
                parsing_stats = {
                    "parser_used": "python-docx",
                    "total_paragraphs": len(document.paragraphs),
                    "total_tables": len(document.tables),
                    "total_sections": len(sections),
                    "total_chunks": len(chunks),
                    "average_chunk_size": sum(len(chunk.content) for chunk in chunks) / len(chunks) if chunks else 0
                }
                
                return ParsedDocument(
                    content=full_content,
                    chunks=chunks,
                    metadata=metadata,
                    sections=sections,
                    parsing_stats=parsing_stats
                )
                
            finally:
                # Clean up temporary file
                os.unlink(temp_file_path)
                
        except Exception as e:
            error_msg = f"DOCX parsing failed for {filename}: {str(e)}"
            logger.error(f"{error_msg}\n{traceback.format_exc()}")
            raise DocumentParsingError(error_msg) from e
    
    def _extract_pdf_metadata(self, pdf_document, filename: str) -> DocumentMetadata:
        """Extract metadata from PDF document."""
        
        metadata_dict = pdf_document.metadata
        
        return DocumentMetadata(
            title=metadata_dict.get('title') or Path(filename).stem,
            author=metadata_dict.get('author'),
            subject=metadata_dict.get('subject'),
            creator=metadata_dict.get('creator'),
            creation_date=metadata_dict.get('creationDate'),
            modification_date=metadata_dict.get('modDate'),
            page_count=len(pdf_document)
        )
    
    def _extract_docx_metadata(self, document, filename: str) -> DocumentMetadata:
        """Extract metadata from DOCX document."""
        
        core_props = document.core_properties
        
        return DocumentMetadata(
            title=core_props.title or Path(filename).stem,
            author=core_props.author,
            subject=core_props.subject,
            creator=core_props.creator,
            creation_date=str(core_props.created) if core_props.created else None,
            modification_date=str(core_props.modified) if core_props.modified else None,
            page_count=1  # DOCX doesn't have clear page concept
        )
    
    def _extract_table_text(self, table) -> str:
        """Extract text from a DOCX table."""
        
        table_text = ""
        for row in table.rows:
            row_text = "\t".join([cell.text.strip() for cell in row.cells])
            if row_text.strip():
                table_text += row_text + "\n"
        
        return table_text
    
    def _create_chunks(
        self, 
        text: str, 
        page_number: int, 
        start_chunk_index: int,
        start_position: int
    ) -> List[DocumentChunk]:
        """Create text chunks from given text."""
        
        chunks = []
        
        # Clean text
        text = re.sub(r'\s+', ' ', text.strip())
        
        if len(text) <= self.chunk_size:
            # Text is small enough for a single chunk
            chunks.append(DocumentChunk(
                content=text,
                page_number=page_number,
                chunk_index=start_chunk_index,
                start_position=start_position,
                end_position=start_position + len(text)
            ))
        else:
            # Split text into overlapping chunks
            chunk_index = start_chunk_index
            pos = 0
            
            while pos < len(text):
                # Calculate chunk end position
                end_pos = min(pos + self.chunk_size, len(text))
                
                # Try to break at word boundary if not at end
                if end_pos < len(text):
                    # Look for space within overlap distance
                    for i in range(end_pos, max(pos + self.chunk_size - self.chunk_overlap, pos), -1):
                        if text[i] == ' ':
                            end_pos = i
                            break
                
                # Extract chunk content
                chunk_content = text[pos:end_pos].strip()
                
                if chunk_content:
                    chunks.append(DocumentChunk(
                        content=chunk_content,
                        page_number=page_number,
                        chunk_index=chunk_index,
                        start_position=start_position + pos,
                        end_position=start_position + end_pos,
                        metadata={
                            "original_length": len(chunk_content),
                            "word_count": len(chunk_content.split())
                        }
                    ))
                    chunk_index += 1
                
                # Move to next chunk with overlap
                pos = max(end_pos - self.chunk_overlap, pos + 1)
        
        return chunks
    
    def _calculate_file_hash(self, file_content: bytes) -> str:
        """Calculate SHA-256 hash of file content for integrity verification."""
        return hashlib.sha256(file_content).hexdigest()
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats."""
        formats = []
        
        if PYMUPDF_AVAILABLE:
            formats.append("pdf")
        
        if PYTHON_DOCX_AVAILABLE:
            formats.append("docx")
        
        return formats
    
    def validate_file_format(self, filename: str) -> bool:
        """Validate if file format is supported."""
        file_extension = Path(filename).suffix.lower()
        return file_extension in self.allowed_extensions
    
    def get_parsing_stats(self) -> Dict[str, Any]:
        """Get current parsing service statistics."""
        return {
            "supported_formats": self.get_supported_formats(),
            "max_file_size_mb": settings.MAX_FILE_SIZE_MB,
            "max_pages": self.max_pages,
            "chunk_size": self.chunk_size,
            "chunk_overlap": self.chunk_overlap,
            "dependencies": {
                "pymupdf_available": PYMUPDF_AVAILABLE,
                "python_docx_available": PYTHON_DOCX_AVAILABLE
            }
        }
    
    async def get_parser_status(self) -> Dict[str, Any]:
        """Get parser status and availability information."""
        return {
            "pdf_available": PYMUPDF_AVAILABLE,
            "docx_available": PYTHON_DOCX_AVAILABLE,
            "yake_enabled": settings.ENABLE_YAKE_SEARCH,
            "supported_formats": self.get_supported_formats(),
            "max_file_size_mb": settings.MAX_FILE_SIZE_MB,
            "max_pages": settings.MAX_PAGES,
            "max_memory_per_session_mb": settings.MAX_MEMORY_PER_SESSION_MB
        }


# Global document parser service instance
document_parser = DocumentParserService()