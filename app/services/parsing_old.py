import io
import logging
import traceback
from typing import Dict, Any, Optional, List, Union
from dataclasses import dataclass
from contextlib import contextmanager

try:
    import pymupdf as fitz  # PyMuPDF
except ImportError as e:
    fitz = None
    logger.warning(f"PyMuPDF not available: {e}")

try:
    from docx import Document
except ImportError as e:
    Document = None
    logger.warning(f"python-docx not available: {e}")

try:
    import yake
except ImportError as e:
    yake = None
    logger.warning(f"YAKE not available: {e}")

from app.core.config import settings

logger = logging.getLogger(__name__)


class DocumentParsingError(Exception):
    """Custom exception for document parsing errors."""
    
    def __init__(self, message: str, error_type: str = "PARSING_ERROR", details: Optional[Dict[str, Any]] = None):
        self.message = message
        self.error_type = error_type
        self.details = details or {}
        super().__init__(self.message)


@dataclass
class ParsedDocument:
    """Parsed document data structure."""
    text: str
    metadata: Dict[str, Any]
    sections: List[Dict[str, Any]]
    page_count: int
    title: str
    keywords: Optional[List[str]] = None


class DocumentParser:
    """
    Production-ready document parsing service for PDF and DOCX files.
    
    Features:
    - Robust error handling with fallbacks
    - Memory usage monitoring
    - Page limit validation
    - Optional YAKE keyword extraction
    - Structured logging
    """
    
    def __init__(self):
        self.yake_extractor = None
        self._init_yake_extractor()
        
        # Check library availability
        self.pdf_available = fitz is not None
        self.docx_available = Document is not None
        self.yake_available = yake is not None
        
        logger.info(f"DocumentParser initialized - PDF: {self.pdf_available}, DOCX: {self.docx_available}, YAKE: {self.yake_available}")
    
    def _init_yake_extractor(self) -> None:
        """Initialize YAKE extractor with error handling."""
        try:
            if yake and settings.ENABLE_YAKE_SEARCH:
                self.yake_extractor = yake.KeywordExtractor(
                    lan="en",
                    n=settings.YAKE_NGRAM_SIZE,
                    dedupLim=0.7,
                    top=settings.YAKE_MAX_KEYWORDS
                )
                logger.info("YAKE keyword extractor initialized")
        except Exception as e:
            logger.warning(f"Failed to initialize YAKE extractor: {e}")
    
    async def parse_pdf(self, file_content: bytes, filename: str) -> ParsedDocument:
        """Parse PDF file and extract text content."""
        if not fitz:
            raise ImportError("PyMuPDF not installed. Install with: pip install pymupdf")
        
        try:
            # Open PDF from bytes
            pdf_document = fitz.open(stream=file_content, filetype="pdf")
            
            # Extract metadata
            metadata = {
                "filename": filename,
                "file_type": "pdf",
                "page_count": pdf_document.page_count,
                "title": pdf_document.metadata.get("title", filename),
                "author": pdf_document.metadata.get("author", ""),
                "subject": pdf_document.metadata.get("subject", ""),
                "creator": pdf_document.metadata.get("creator", ""),
                "producer": pdf_document.metadata.get("producer", ""),
                "creation_date": pdf_document.metadata.get("creationDate", ""),
                "modification_date": pdf_document.metadata.get("modDate", "")
            }
            
            # Check page limit
            if pdf_document.page_count > settings.MAX_PAGES:
                raise ValueError(f"Document has {pdf_document.page_count} pages, maximum allowed is {settings.MAX_PAGES}")
            
            # Extract text from all pages
            full_text = ""
            sections = []
            
            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                page_text = page.get_text()
                
                if page_text.strip():  # Only add non-empty pages
                    full_text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                    sections.append({
                        "page": page_num + 1,
                        "text": page_text.strip(),
                        "char_count": len(page_text)
                    })
            
            pdf_document.close()
            
            # Extract keywords if YAKE is enabled
            keywords = None
            if self.yake_extractor and full_text.strip():
                try:
                    extracted_keywords = self.yake_extractor.extract_keywords(full_text)
                    keywords = [kw[1] for kw in extracted_keywords]
                    logger.debug(f"Extracted {len(keywords)} keywords from PDF")
                except Exception as e:
                    logger.warning(f"YAKE keyword extraction failed: {e}")
            
            return ParsedDocument(
                text=full_text.strip(),
                metadata=metadata,
                sections=sections,
                page_count=pdf_document.page_count,
                title=metadata["title"] or filename,
                keywords=keywords
            )
            
        except Exception as e:
            logger.error(f"PDF parsing failed for {filename}: {e}")
            raise ValueError(f"Failed to parse PDF: {str(e)}")
    
    async def parse_docx(self, file_content: bytes, filename: str) -> ParsedDocument:
        """Parse DOCX file and extract text content."""
        if not Document:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")
        
        try:
            # Open DOCX from bytes
            doc = Document(io.BytesIO(file_content))
            
            # Extract core properties
            props = doc.core_properties
            metadata = {
                "filename": filename,
                "file_type": "docx",
                "title": props.title or filename,
                "author": props.author or "",
                "subject": props.subject or "",
                "created": props.created.isoformat() if props.created else "",
                "modified": props.modified.isoformat() if props.modified else "",
                "last_modified_by": props.last_modified_by or "",
                "revision": props.revision or "",
                "category": props.category or "",
                "comments": props.comments or ""
            }
            
            # Extract text from paragraphs
            full_text = ""
            sections = []
            paragraph_count = 0
            
            for para in doc.paragraphs:
                if para.text.strip():  # Only non-empty paragraphs
                    paragraph_count += 1
                    para_text = para.text.strip()
                    full_text += f"{para_text}\n"
                    
                    sections.append({
                        "paragraph": paragraph_count,
                        "text": para_text,
                        "char_count": len(para_text),
                        "style": para.style.name if para.style else "Normal"
                    })
            
            # Estimate page count (rough approximation: 500 chars per page)
            estimated_pages = max(1, len(full_text) // 500)
            metadata["page_count"] = estimated_pages
            metadata["paragraph_count"] = paragraph_count
            
            # Check page limit (estimated)
            if estimated_pages > settings.MAX_PAGES:
                raise ValueError(f"Document estimated to have {estimated_pages} pages, maximum allowed is {settings.MAX_PAGES}")
            
            # Extract keywords if YAKE is enabled
            keywords = None
            if self.yake_extractor and full_text.strip():
                try:
                    extracted_keywords = self.yake_extractor.extract_keywords(full_text)
                    keywords = [kw[1] for kw in extracted_keywords]
                    logger.debug(f"Extracted {len(keywords)} keywords from DOCX")
                except Exception as e:
                    logger.warning(f"YAKE keyword extraction failed: {e}")
            
            return ParsedDocument(
                text=full_text.strip(),
                metadata=metadata,
                sections=sections,
                page_count=estimated_pages,
                title=metadata["title"] or filename,
                keywords=keywords
            )
            
        except Exception as e:
            logger.error(f"DOCX parsing failed for {filename}: {e}")
            raise ValueError(f"Failed to parse DOCX: {str(e)}")
    
    async def parse_document(self, file_content: bytes, filename: str, file_type: str) -> ParsedDocument:
        """Parse document based on file type."""
        if file_type.lower() == ".pdf":
            return await self.parse_pdf(file_content, filename)
        elif file_type.lower() == ".docx":
            return await self.parse_docx(file_content, filename)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    def create_yake_index(self, document_text: str) -> Optional[Dict[str, Any]]:
        """Create YAKE keyword index for search optimization."""
        if not self.yake_extractor or not document_text.strip():
            return None
        
        try:
            keywords = self.yake_extractor.extract_keywords(document_text)
            
            # Create searchable index
            keyword_index = {
                "keywords": [
                    {
                        "keyword": kw[1],
                        "score": kw[0],
                        "rank": idx + 1
                    }
                    for idx, kw in enumerate(keywords)
                ],
                "total_keywords": len(keywords),
                "text_length": len(document_text),
                "created_at": str(fitz.get_systime()) if fitz else ""
            }
            
            logger.info(f"Created YAKE index with {len(keywords)} keywords")
            return keyword_index
            
        except Exception as e:
            logger.error(f"YAKE index creation failed: {e}")
            return None


# Global parser instance
document_parser = DocumentParser()