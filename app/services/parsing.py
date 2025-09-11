import io
import logging
import traceback
from typing import Dict, Any, Optional, List, Union
from dataclasses import dataclass
from contextlib import contextmanager

try:
    import pymupdf as fitz  # PyMuPDF
except ImportError as e:
    fitz = None
    logging.warning(f"PyMuPDF not available: {e}")

try:
    from docx import Document
except ImportError as e:
    Document = None
    logging.warning(f"python-docx not available: {e}")

try:
    import yake
except ImportError as e:
    yake = None
    logging.warning(f"YAKE not available: {e}")

from app.core.config import settings

logger = logging.getLogger(__name__)


class DocumentParsingError(Exception):
    """Custom exception for document parsing errors."""
    
    def __init__(self, message: str, error_type: str = "PARSING_ERROR", details: Optional[Dict[str, Any]] = None):
        self.message = message
        self.error_type = error_type
        self.details = details or {}
        super().__init__(self.message)


@dataclass
class ParsedDocument:
    """Parsed document data structure."""
    text: str
    metadata: Dict[str, Any]
    sections: List[Dict[str, Any]]
    page_count: int
    title: str
    keywords: Optional[List[str]] = None
    
    def get_memory_estimate_mb(self) -> float:
        """Estimate memory usage in MB."""
        try:
            size = len(self.text.encode('utf-8'))
            size += len(str(self.metadata).encode('utf-8'))
            size += len(str(self.sections).encode('utf-8'))
            if self.keywords:
                size += len(str(self.keywords).encode('utf-8'))
            return size / (1024 * 1024)
        except Exception:
            return 0.0


class DocumentParser:
    """
    Production-ready document parsing service for PDF and DOCX files.
    
    Features:
    - Robust error handling with fallbacks
    - Memory usage monitoring
    - Page limit validation
    - Optional YAKE keyword extraction
    - Structured logging
    """
    
    def __init__(self):
        self.yake_extractor = None
        self._init_yake_extractor()
        
        # Check library availability
        self.pdf_available = fitz is not None
        self.docx_available = Document is not None
        self.yake_available = yake is not None
        
        logger.info(f"DocumentParser initialized - PDF: {self.pdf_available}, DOCX: {self.docx_available}, YAKE: {self.yake_available}")
    
    def _init_yake_extractor(self) -> None:
        """Initialize YAKE extractor with error handling."""
        try:
            if yake and settings.ENABLE_YAKE_SEARCH:
                self.yake_extractor = yake.KeywordExtractor(
                    lan="en",
                    n=settings.YAKE_NGRAM_SIZE,
                    dedupLim=0.7,
                    top=settings.YAKE_MAX_KEYWORDS
                )
                logger.info("YAKE keyword extractor initialized")
        except Exception as e:
            logger.warning(f"Failed to initialize YAKE extractor: {e}")
    
    @contextmanager
    def _safe_pdf_context(self, file_content: bytes, filename: str):
        """Context manager for safe PDF handling with automatic cleanup."""
        pdf_document = None
        try:
            pdf_document = fitz.open(stream=file_content, filetype="pdf")
            yield pdf_document
        except Exception as e:
            logger.error(f"PDF context error for {filename}: {e}")
            raise DocumentParsingError(
                f"Failed to open PDF: {str(e)}",
                error_type="PDF_CORRUPT",
                details={"filename": filename, "error": str(e)}
            )
        finally:
            if pdf_document:
                try:
                    pdf_document.close()
                    logger.debug(f"PDF document closed for {filename}")
                except Exception as e:
                    logger.warning(f"Failed to close PDF {filename}: {e}")
    
    async def parse_pdf(self, file_content: bytes, filename: str) -> ParsedDocument:
        """
        Parse PDF file and extract text content with robust error handling.
        
        Args:
            file_content: PDF file content as bytes
            filename: Original filename for logging/metadata
            
        Returns:
            ParsedDocument with extracted text and metadata
            
        Raises:
            DocumentParsingError: If parsing fails with details
        """
        if not self.pdf_available:
            raise DocumentParsingError(
                "PyMuPDF not available. Install with: pip install pymupdf",
                error_type="LIBRARY_MISSING",
                details={"library": "pymupdf", "filename": filename}
            )
        
        try:
            logger.info(f"Starting PDF parsing for {filename} ({len(file_content)} bytes)")
            
            # Validate file size
            file_size_mb = len(file_content) / (1024 * 1024)
            if file_size_mb > settings.MAX_FILE_SIZE_MB:
                raise DocumentParsingError(
                    f"PDF file too large: {file_size_mb:.2f}MB > {settings.MAX_FILE_SIZE_MB}MB",
                    error_type="FILE_TOO_LARGE",
                    details={"size_mb": file_size_mb, "max_size_mb": settings.MAX_FILE_SIZE_MB}
                )
            
            with self._safe_pdf_context(file_content, filename) as pdf_document:
                # Check page count
                if pdf_document.page_count > settings.MAX_PAGES:
                    raise DocumentParsingError(
                        f"PDF has too many pages: {pdf_document.page_count} > {settings.MAX_PAGES}",
                        error_type="TOO_MANY_PAGES",
                        details={"pages": pdf_document.page_count, "max_pages": settings.MAX_PAGES}
                    )
                
                # Extract metadata safely
                metadata = self._extract_pdf_metadata(pdf_document, filename)
                
                # Extract text from all pages
                full_text, sections = self._extract_pdf_text(pdf_document)
                
                # Validate extracted content
                if not full_text.strip():
                    logger.warning(f"No text extracted from PDF {filename}")
                    raise DocumentParsingError(
                        "No readable text found in PDF",
                        error_type="NO_TEXT_CONTENT",
                        details={"filename": filename, "pages": pdf_document.page_count}
                    )
                
                # Extract keywords if YAKE is enabled
                keywords = self._extract_keywords_safe(full_text)
                
                parsed_doc = ParsedDocument(
                    text=full_text.strip(),
                    metadata=metadata,
                    sections=sections,
                    page_count=pdf_document.page_count,
                    title=metadata.get("title", filename),
                    keywords=keywords
                )
                
                # Check memory usage
                memory_mb = parsed_doc.get_memory_estimate_mb()
                if memory_mb > settings.MAX_MEMORY_PER_SESSION_MB:
                    logger.warning(f"PDF {filename} uses {memory_mb:.2f}MB memory")
                
                logger.info(f"Successfully parsed PDF {filename}: {pdf_document.page_count} pages, {len(full_text)} chars, {memory_mb:.2f}MB")
                return parsed_doc
                
        except DocumentParsingError:
            raise
        except Exception as e:
            logger.error(f"Unexpected error parsing PDF {filename}: {e}\n{traceback.format_exc()}")
            raise DocumentParsingError(
                f"Unexpected PDF parsing error: {str(e)}",
                error_type="UNEXPECTED_ERROR",
                details={"filename": filename, "error": str(e), "traceback": traceback.format_exc()}
            )
    
    def _extract_pdf_metadata(self, pdf_document, filename: str) -> Dict[str, Any]:
        """Safely extract PDF metadata."""
        try:
            metadata = {
                "filename": filename,
                "file_type": "pdf",
                "page_count": pdf_document.page_count,
                "title": pdf_document.metadata.get("title", filename) or filename,
                "author": pdf_document.metadata.get("author", ""),
                "subject": pdf_document.metadata.get("subject", ""),
                "creator": pdf_document.metadata.get("creator", ""),
                "producer": pdf_document.metadata.get("producer", ""),
                "creation_date": pdf_document.metadata.get("creationDate", ""),
                "modification_date": pdf_document.metadata.get("modDate", "")
            }
            return metadata
        except Exception as e:
            logger.warning(f"Failed to extract PDF metadata from {filename}: {e}")
            return {
                "filename": filename,
                "file_type": "pdf",
                "page_count": getattr(pdf_document, 'page_count', 0),
                "title": filename,
                "author": "",
                "subject": "",
                "creator": "",
                "producer": "",
                "creation_date": "",
                "modification_date": ""
            }
    
    def _extract_pdf_text(self, pdf_document) -> tuple[str, List[Dict[str, Any]]]:
        """Extract text from PDF with error handling per page."""
        full_text = ""
        sections = []
        successful_pages = 0
        
        for page_num in range(pdf_document.page_count):
            try:
                page = pdf_document[page_num]
                page_text = page.get_text()
                
                if page_text.strip():  # Only add non-empty pages
                    full_text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                    sections.append({
                        "page": page_num + 1,
                        "text": page_text.strip(),
                        "char_count": len(page_text),
                        "status": "success"
                    })
                    successful_pages += 1
                else:
                    sections.append({
                        "page": page_num + 1,
                        "text": "",
                        "char_count": 0,
                        "status": "empty"
                    })
                    
            except Exception as e:
                logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                sections.append({
                    "page": page_num + 1,
                    "text": "",
                    "char_count": 0,
                    "status": "error",
                    "error": str(e)
                })
        
        logger.info(f"Extracted text from {successful_pages}/{pdf_document.page_count} pages")
        return full_text, sections
    
    async def parse_docx(self, file_content: bytes, filename: str) -> ParsedDocument:
        """
        Parse DOCX file and extract text content with robust error handling.
        """
        if not self.docx_available:
            raise DocumentParsingError(
                "python-docx not available. Install with: pip install python-docx",
                error_type="LIBRARY_MISSING",
                details={"library": "python-docx", "filename": filename}
            )
        
        try:
            logger.info(f"Starting DOCX parsing for {filename} ({len(file_content)} bytes)")
            
            # Validate file size
            file_size_mb = len(file_content) / (1024 * 1024)
            if file_size_mb > settings.MAX_FILE_SIZE_MB:
                raise DocumentParsingError(
                    f"DOCX file too large: {file_size_mb:.2f}MB > {settings.MAX_FILE_SIZE_MB}MB",
                    error_type="FILE_TOO_LARGE",
                    details={"size_mb": file_size_mb, "max_size_mb": settings.MAX_FILE_SIZE_MB}
                )
            
            # Open DOCX from bytes with error handling
            try:
                doc = Document(io.BytesIO(file_content))
            except Exception as e:
                raise DocumentParsingError(
                    f"Failed to open DOCX file: {str(e)}",
                    error_type="DOCX_CORRUPT",
                    details={"filename": filename, "error": str(e)}
                )
            
            # Extract metadata safely
            metadata = self._extract_docx_metadata(doc, filename)
            
            # Extract text from paragraphs
            full_text, sections = self._extract_docx_text(doc)
            
            # Validate extracted content
            if not full_text.strip():
                logger.warning(f"No text extracted from DOCX {filename}")
                raise DocumentParsingError(
                    "No readable text found in DOCX",
                    error_type="NO_TEXT_CONTENT",
                    details={"filename": filename}
                )
            
            # Estimate page count (rough approximation: 500 chars per page)
            estimated_pages = max(1, len(full_text) // 500)
            metadata["page_count"] = estimated_pages
            
            # Check page limit (estimated)
            if estimated_pages > settings.MAX_PAGES:
                raise DocumentParsingError(
                    f"DOCX estimated to have too many pages: {estimated_pages} > {settings.MAX_PAGES}",
                    error_type="TOO_MANY_PAGES",
                    details={"estimated_pages": estimated_pages, "max_pages": settings.MAX_PAGES}
                )
            
            # Extract keywords if YAKE is enabled
            keywords = self._extract_keywords_safe(full_text)
            
            parsed_doc = ParsedDocument(
                text=full_text.strip(),
                metadata=metadata,
                sections=sections,
                page_count=estimated_pages,
                title=metadata.get("title", filename) or filename,
                keywords=keywords
            )
            
            # Check memory usage
            memory_mb = parsed_doc.get_memory_estimate_mb()
            logger.info(f"Successfully parsed DOCX {filename}: ~{estimated_pages} pages, {len(full_text)} chars, {memory_mb:.2f}MB")
            return parsed_doc
            
        except DocumentParsingError:
            raise
        except Exception as e:
            logger.error(f"Unexpected error parsing DOCX {filename}: {e}\n{traceback.format_exc()}")
            raise DocumentParsingError(
                f"Unexpected DOCX parsing error: {str(e)}",
                error_type="UNEXPECTED_ERROR",
                details={"filename": filename, "error": str(e), "traceback": traceback.format_exc()}
            )
    
    def _extract_docx_metadata(self, doc, filename: str) -> Dict[str, Any]:
        """Safely extract DOCX metadata."""
        try:
            props = doc.core_properties
            metadata = {
                "filename": filename,
                "file_type": "docx",
                "title": props.title or filename,
                "author": props.author or "",
                "subject": props.subject or "",
                "created": props.created.isoformat() if props.created else "",
                "modified": props.modified.isoformat() if props.modified else "",
                "last_modified_by": props.last_modified_by or "",
                "revision": props.revision or "",
                "category": props.category or "",
                "comments": props.comments or ""
            }
            return metadata
        except Exception as e:
            logger.warning(f"Failed to extract DOCX metadata from {filename}: {e}")
            return {
                "filename": filename,
                "file_type": "docx",
                "title": filename,
                "author": "",
                "subject": "",
                "created": "",
                "modified": "",
                "last_modified_by": "",
                "revision": "",
                "category": "",
                "comments": ""
            }
    
    def _extract_docx_text(self, doc) -> tuple[str, List[Dict[str, Any]]]:
        """Extract text from DOCX paragraphs with error handling."""
        full_text = ""
        sections = []
        paragraph_count = 0
        successful_paragraphs = 0
        
        for para in doc.paragraphs:
            try:
                if para.text.strip():  # Only non-empty paragraphs
                    paragraph_count += 1
                    para_text = para.text.strip()
                    full_text += f"{para_text}\n"
                    
                    sections.append({
                        "paragraph": paragraph_count,
                        "text": para_text,
                        "char_count": len(para_text),
                        "style": para.style.name if para.style else "Normal",
                        "status": "success"
                    })
                    successful_paragraphs += 1
                    
            except Exception as e:
                logger.warning(f"Failed to extract paragraph {paragraph_count + 1}: {e}")
                sections.append({
                    "paragraph": paragraph_count + 1,
                    "text": "",
                    "char_count": 0,
                    "style": "Unknown",
                    "status": "error",
                    "error": str(e)
                })
        
        logger.info(f"Extracted text from {successful_paragraphs} paragraphs")
        return full_text, sections
    
    def _extract_keywords_safe(self, text: str) -> Optional[List[str]]:
        """Safely extract keywords with YAKE."""
        if not self.yake_extractor or not text.strip():
            return None
        
        try:
            extracted_keywords = self.yake_extractor.extract_keywords(text)
            keywords = [kw[1] for kw in extracted_keywords]
            logger.debug(f"Extracted {len(keywords)} keywords")
            return keywords
        except Exception as e:
            logger.warning(f"YAKE keyword extraction failed: {e}")
            return None
    
    async def parse_document(self, file_content: bytes, filename: str, file_type: str) -> ParsedDocument:
        """
        Parse document based on file type with comprehensive error handling.
        
        Args:
            file_content: File content as bytes
            filename: Original filename
            file_type: File extension (.pdf or .docx)
            
        Returns:
            ParsedDocument with extracted content
            
        Raises:
            DocumentParsingError: If parsing fails
        """
        try:
            file_type_lower = file_type.lower()
            
            if file_type_lower == ".pdf":
                return await self.parse_pdf(file_content, filename)
            elif file_type_lower == ".docx":
                return await self.parse_docx(file_content, filename)
            else:
                raise DocumentParsingError(
                    f"Unsupported file type: {file_type}",
                    error_type="UNSUPPORTED_FORMAT",
                    details={"file_type": file_type, "supported": [".pdf", ".docx"]}
                )
                
        except DocumentParsingError:
            raise
        except Exception as e:
            logger.error(f"Unexpected error in parse_document: {e}")
            raise DocumentParsingError(
                f"Unexpected parsing error: {str(e)}",
                error_type="UNEXPECTED_ERROR",
                details={"filename": filename, "file_type": file_type, "error": str(e)}
            )
    
    def get_parser_status(self) -> Dict[str, Any]:
        """Get parser status and availability information."""
        return {
            "pdf_available": self.pdf_available,
            "docx_available": self.docx_available,
            "yake_available": self.yake_available,
            "yake_enabled": settings.ENABLE_YAKE_SEARCH,
            "supported_formats": [".pdf", ".docx"] if self.pdf_available or self.docx_available else [],
            "max_file_size_mb": settings.MAX_FILE_SIZE_MB,
            "max_pages": settings.MAX_PAGES,
            "max_memory_per_session_mb": settings.MAX_MEMORY_PER_SESSION_MB
        }


# Global parser instance
document_parser = DocumentParser()